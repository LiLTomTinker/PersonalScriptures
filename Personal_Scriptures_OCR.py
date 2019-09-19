'''
This program takes a pdf version of a user's patriarchal blessing from lds.org and creates a formatted draft of
it as a Word document. The user then needs to scan through the text themselves and fix any OCR errors.
'''

from docx import Document
from docx.shared import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from wand.image import Image
from PIL import Image as PI
import pyocr
import pyocr.builders
import io
import os

tool = pyocr.get_available_tools()[0]
lang = tool.get_available_languages()[1]

req_image = []
final_text = []

user_document = input("Input document name in quotes (include .pdf): ")

image_pdf = Image(filename=user_document, resolution=230)
image_jpeg = image_pdf.convert('JPEG')

print image_jpeg.size
i = 1
first_page = image_jpeg.sequence[0]

img_page = Image(image=image_jpeg.sequence[0])
regular_page_height = int(0.9122734 * img_page.height)
img_page.crop(int(0.0515909 * img_page.width), int(0.2575078 * img_page.height), width=int(0.8957759 * img_page.width), height=int(0.6813658 * img_page.height))
img_page.save(filename='PB-raw0.jpeg')
top = Image(filename='PB-raw0.jpeg')

for img in image_jpeg.sequence[1:]:
    img_page = Image(image=img)
    #img_page.crop(95,145,1845,2355) #LTRB
    img_page.crop(int(0.0515909 * img_page.width), int(0.05031 * img_page.height), width=int(0.8957759 * img_page.width), height=int(0.9122734 * img_page.height))
    #req_image.append(img_page.make_blob('JPEG'))
#print image_jpeg.crop(10, 20, width=45, height=220)
    img_page.save(filename='PB-raw' + str(i) + '.jpeg')
    i += 1

#with Image(filename='PB-raw0.jpeg') as top:
with Image(width=top.width,
           height=(top.height + (regular_page_height * (i-1))))  as stitch:
           stitch.composite(image=top, left=0, top=0)
           stitch.composite(image=Image(filename='PB-raw1.jpeg'), left=0, top=top.height)
           for j in range(2,i):
               stitch.composite(image=Image(filename='PB-raw' + str(j) + '.jpeg'), left=0, top=top.height + (regular_page_height * (j-1)))
           stitch.save(filename='PB-raw-stitch.jpeg')  #with Image(filename='PB-raw' + str(j) + '.jpeg') as bottom:


for k in range(0,i):
    os.remove('PB-raw' + str(k) + '.jpeg')

# Need to crop page one header information and then header entirely.

raw_stitch = Image(filename='PB-raw-stitch.jpeg').make_blob('JPEG')
txt = tool.image_to_string(
    PI.open(io.BytesIO(raw_stitch)),
    lang=lang,
    builder=pyocr.builders.TextBuilder()
)
txt = txt.replace('\n',' ')
txt = txt.replace('  ','\n')
for i in range(0,10):
    txt = txt.replace('\n\n','\n')
txt = txt.replace('\n','\n\n')
txt = txt.replace('&','a')
txt = txt.replace('1','I')

clean_up = 'bcdefghijklmnopqqrstuvwxyzBCDEFGHJKLMNOPQRSTUVWXYZ'
for i in range(0,len(clean_up)):
    txt = txt.replace(' ' + clean_up[i] + ' ', ' ' + clean_up[i])
    txt = txt.replace('.' + clean_up[i], '@' + clean_up[i])
    txt = txt.replace(clean_up[i] + ']', clean_up[i] + 'l')
    txt = txt.replace('a]','al')
    txt = txt.replace('a[','al')
    txt = txt.replace(clean_up[i] + '[', clean_up[i] + 'l')

txt = txt.replace(' ] ', ' I ')
txt = txt.replace(' [ ', ' I ')

numParagraphs = txt.count('\n\n') + 1

txt_list = txt.split('\n\n')
#final_text.append(txt

ocr_raw = Document()
ocr_raw.add_paragraph('')
ocr_raw.paragraphs[0].add_run(txt) # Each txt is generated per a page from the PDF
ocr_raw.save('OCR-raw.docx')
user_input = input("Type 'p' for PARAGRAPHS or 'c' for CHAPTERS (include quotes): ")
document = Document('OCR-raw.docx')
user_input = str(user_input)
if user_input == 'p':
    section_string = "PARAGRAPH "
else:
    section_string = "CHAPTER "


new_document = Document('PB-Template.docx')
#numParagraphs = len(document.paragraphs)
numTemplateParagraphs = len(new_document.paragraphs)

header_name_size = 40
drop_cap_font_size = 46
body_text_size = 11

#new_document.paragraphs[2].text = document.paragraphs[0].text[0]

for i in range(0,3):
    new_document.paragraphs[i].alignment = WD_ALIGN_PARAGRAPH.CENTER

section = new_document.sections[0] # <--
sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'),'1')


for j in range(0, numParagraphs):
    if j > 1: new_document.add_paragraph()
    #document.paragraphs[j].text = document.paragraphs[j].text.replace('\n', '')
    txt_list[j] = txt_list[j].replace('.', '.\n')
     #we don't want to put a verse number after the last period in the paragraph
    verses_list = txt_list[j].splitlines()
    #if len(verses_list) > 1:
    if j == 0:
        description = new_document.paragraphs[-2].insert_paragraph_before("Write a description ending with a period." + "\n") # (j/2)+1 accounts for the "paragraphs" that are just a newline.

        paragraph_number = description.insert_paragraph_before(section_string + str(j+1) + "\n") # <-- Every paragraph with text will be an even-numbered paragraph
        paragraph_number.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #paragraph_number.runs[0].font.name = "Palatino"
        description_font = description.runs[0].font
        description_font.italic = True
        #description_font.name = "Palatino"

        paragraph = new_document.paragraphs[-1]
        first_space_index = verses_list[0].find(" ")
        first_comma_index = verses_list[0].find(",")
        brother_index = verses_list[0].find("Brother ")
        sister_index = verses_list[0].find("Sister ")
        if sister_index > 0:
            new_document.paragraphs[2].text = verses_list[0].strip('\n').upper()[sister_index + len("Sister "):first_comma_index]
            new_document.paragraphs[2].runs[0].font.size = Pt(header_name_size) # paragraphs[2] is the second header line
            #new_document.paragraphs[2].runs[0].font.name = "Palatino"
        if brother_index > 0:
            new_document.paragraphs[2].text = verses_list[0].strip('\n').upper()[brother_index + len("Brother "):first_comma_index]
            new_document.paragraphs[2].runs[0].font.size = Pt(header_name_size)
            #new_document.paragraphs[2].runs[0].font.name = "Palatino"

        first_word_letters = paragraph.add_run(verses_list[0].strip('\n')[1:first_space_index].upper())
        paragraph.add_run(verses_list[0].strip('\n')[first_space_index:])
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        drop_cap = new_document.paragraphs[numTemplateParagraphs].runs[0]
        drop_cap.text = verses_list[0].strip('\n')[:1]
        drop_cap.font.size = Pt(drop_cap_font_size)
        new_document.paragraphs[numTemplateParagraphs].runs[0]
        for i in range(1,len(verses_list)):
            paragraph = new_document.add_paragraph("  " + str(i + 1) + verses_list[i].strip('\n'))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in paragraph.runs:
            #run.font.name = "Palatino"
            run.font.size = Pt(11)



    else:
        if j == 1: new_document.add_paragraph()
        paragraph_number = new_document.add_paragraph(section_string + str(j+1) + "\n") # (j/2)+1 accounts for the "paragraphs" that are just a newline.
        paragraph_number.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #paragraph_number.runs[0].font.name = "Palatino"
        description = new_document.add_paragraph("Write a description ending with a period." + "\n")
        description_font = description.runs[0].font
        description_font.italic = True
        #description_font.name = "Palatino"

        paragraph = new_document.add_paragraph()
        first_letter = paragraph.add_run(verses_list[0].strip('\n')[:1].upper())
        first_space_index = verses_list[0].strip('\n').find(" ")
        first_word_letters = paragraph.add_run(verses_list[0].strip('\n')[1:first_space_index].upper())
        paragraph.add_run(verses_list[0].strip('\n')[first_space_index:])
        for run in paragraph.runs:
            #run.font.name = 'Palatino'
            run.font.size = Pt(body_text_size)
        first_word_letters.font.size = Pt(9.5)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for i in range(1,len(verses_list)):
            paragraph = new_document.add_paragraph("  " + str(i + 1) + verses_list[i].strip('\n'))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in paragraph.runs:
                font = run.font
                #font.name = 'Palatino'
                font.size = Pt(body_text_size)
#new_document.paragraphs[j].insert_paragraph_before("PARAGRAPH " + str((j / 2) + 1) + "\n\n")

paragraph_format = new_document.styles['Normal'].paragraph_format
paragraphs = new_document.paragraphs

paragraph_format.line_spacing = 1
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)

#paragraphs[numTemplateParagraphs].runs[0].font.size = Pt(48)


section = new_document.sections[2] # <-- This is supposed to be everything but the header

sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'),'2')


sections = new_document.sections
for section in sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

section = document.sections[0]
section.page_height = Mm(178)
section.page_width = Mm(127)


paragraphs = new_document.paragraphs
for paragraph in paragraphs:
    for run in paragraph.runs:
        run.font.name = "Palatino"


new_document.save(user_document[:-4] + '_formatted' + '.docx')
print("The raw text of your patriarchal blessing has been formatted. Open in Word. Enjoy!")
