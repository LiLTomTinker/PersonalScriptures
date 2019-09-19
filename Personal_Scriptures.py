'''
This program takes in a Word document as input. The input document is intended to be an unformatted Word document that contains the
raw text of a user's patriarchal blessing, seperated into paragrpahs. The output document is a "scripture format" version of
the user's patriarchal blessing saved to a new Word document with the file name the same as the input document, but appended with
"_formatted"

'''
from docx import Document
from docx.shared import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

user_document = input("Input document name in quotes (include .docx): ")
user_input = input("Type 'p' for PARAGRAPHS, 'c' for CHAPTERS, or 's' for SECTIONS (include quotes): ")
document = Document(user_document)
user_input = str(user_input)
if user_input == 'p':
    section_string = "PARAGRAPH "
elif user_input == 'c':
    section_string = "CHAPTER "
elif user_input == 's':
    section_string = "SECTION "



new_document = Document('PB-Template.docx')
numParagraphs = len(document.paragraphs)
numTemplateParagraphs = len(new_document.paragraphs)

header_name_size = 40
drop_cap_font_size = 46
body_text_size = 11

#new_document.paragraphs[2].text = document.paragraphs[0].text[0]

for i in range(0,3):
    new_document.paragraphs[i].alignment = WD_ALIGN_PARAGRAPH.CENTER

section = new_document.sections[0] # <--
sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'),'1')

for j in range(0, numParagraphs):
    if j > 1: new_document.add_paragraph()
    #document.paragraphs[j].text = document.paragraphs[j].text.replace('\n', '')
    document.paragraphs[j].text = document.paragraphs[j].text.replace('.', '.\n')
     #we don't want to put a verse number after the last period in the paragraph
    verses_list = document.paragraphs[j].text.splitlines()
    if len(verses_list) > 1:
        if ((j/2)+1) == 1:
            description = new_document.paragraphs[-2].insert_paragraph_before("Write a description ending with a period." + "\n") # (j/2)+1 accounts for the "paragraphs" that are just a newline.

            paragraph_number = description.insert_paragraph_before(section_string + str((j / 2) + 1) + "\n") # <-- Every paragraph with text will be an even-numbered paragraph
            paragraph_number.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #paragraph_number.runs[0].font.name = "Palatino"
            description_font = description.runs[0].font
            description_font.italic = True
            #description_font.name = "Palatino"

            paragraph = new_document.paragraphs[-1]
            first_space_index = verses_list[0].find(" ")
            first_comma_index = verses_list[0].find(",")
            brother_index = verses_list[0].find("Brother ")
            sister_index = verses_list[0].find("Sister ")
            if sister_index > 0:
                new_document.paragraphs[2].text = verses_list[0].strip('\n').upper()[sister_index + len("Sister "):first_comma_index]
                new_document.paragraphs[2].runs[0].font.size = Pt(header_name_size) # paragraphs[2] is the second header line
                #new_document.paragraphs[2].runs[0].font.name = "Palatino"
            if brother_index > 0:
                new_document.paragraphs[2].text = verses_list[0].strip('\n').upper()[brother_index + len("Brother "):first_comma_index]
                new_document.paragraphs[2].runs[0].font.size = Pt(header_name_size)
                #new_document.paragraphs[2].runs[0].font.name = "Palatino"

            first_word_letters = paragraph.add_run(verses_list[0].strip('\n')[1:first_space_index].upper())
            paragraph.add_run(verses_list[0].strip('\n')[first_space_index:])
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            drop_cap = new_document.paragraphs[numTemplateParagraphs].runs[0]
            drop_cap.text = verses_list[0].strip('\n')[:1]
            drop_cap.font.size = Pt(drop_cap_font_size)
            new_document.paragraphs[numTemplateParagraphs].runs[0]
            for i in range(1,len(verses_list)):
                paragraph = new_document.add_paragraph("  " + str(i + 1) + verses_list[i].strip('\n'))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in paragraph.runs:
                #run.font.name = "Palatino"
                run.font.size = Pt(11)

        else:
            if ((j/2)+1) == 2: new_document.add_paragraph()
            paragraph_number = new_document.add_paragraph(section_string + str((j / 2) + 1) + "\n") # (j/2)+1 accounts for the "paragraphs" that are just a newline.
            paragraph_number.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #paragraph_number.runs[0].font.name = "Palatino"
            description = new_document.add_paragraph("Write a description ending with a period." + "\n")
            description_font = description.runs[0].font
            description_font.italic = True

            paragraph = new_document.add_paragraph()
            first_letter = paragraph.add_run(verses_list[0].strip('\n')[:1].upper())
            first_space_index = verses_list[0].strip('\n').find(" ")
            first_word_letters = paragraph.add_run(verses_list[0].strip('\n')[1:first_space_index].upper())
            paragraph.add_run(verses_list[0].strip('\n')[first_space_index:])
            for run in paragraph.runs:
                #run.font.name = 'Palatino'
                run.font.size = Pt(body_text_size)
            first_word_letters.font.size = Pt(9.5)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            for i in range(1,len(verses_list)):
                paragraph = new_document.add_paragraph("  " + str(i + 1) + verses_list[i].strip('\n'))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in paragraph.runs:
                    font = run.font
                    #font.name = 'Palatino'
                    font.size = Pt(body_text_size)

paragraph_format = new_document.styles['Normal'].paragraph_format
paragraphs = new_document.paragraphs

paragraph_format.line_spacing = 1
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)

section = new_document.sections[2] # <-- This is supposed to be everything but the header

sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'),'2')

sections = new_document.sections
for section in sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

section = document.sections[0]
section.page_height = Mm(178)
section.page_width = Mm(127)

paragraphs = new_document.paragraphs
for paragraph in paragraphs:
    for run in paragraph.runs:
        run.font.name = "Palatino"

new_document.save(user_document[:-5] + '_formatted' + '.docx')
print("The raw text of your patriarchal blessing has been formatted. Open in Word. Enjoy!")
