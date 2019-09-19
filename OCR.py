from docx import Document
from docx.shared import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement


from wand.image import Image
from PIL import Image as PI
import pyocr
import pyocr.builders
import io
import os

tool = pyocr.get_available_tools()[0]
lang = tool.get_available_languages()[1]

req_image = []
final_text = []

user_document = input("Input document name in quotes (include .pdf): ")

image_pdf = Image(filename=user_document, resolution=230)
image_jpeg = image_pdf.convert('JPEG')

print image_jpeg.size
i = 1
first_page = image_jpeg.sequence[0]

#Crop first page to get first name, date, patriarch name, etc. ?

#FIGURE OUT CROP ISSUES    ----   Creates a new paragraph at each new line on the PDF...
img_page = Image(image=image_jpeg.sequence[0])
regular_page_height = int(0.9122734 * img_page.height)
img_page.crop(int(0.0525909 * img_page.width), int(0.2575078 * img_page.height), width=int(0.8957759 * img_page.width), height=int(0.6813658 * img_page.height))
img_page.save(filename='PB-raw0.jpeg')
top = Image(filename='PB-raw0.jpeg')

for img in image_jpeg.sequence[1:]:
    img_page = Image(image=img)
    #img_page.crop(95,145,1845,2355) #LTRB
    img_page.crop(int(0.0525909 * img_page.width), int(0.05031 * img_page.height), width=int(0.8957759 * img_page.width), height=int(0.9122734 * img_page.height))
    #req_image.append(img_page.make_blob('JPEG'))
#print image_jpeg.crop(10, 20, width=45, height=220)
    img_page.save(filename='PB-raw' + str(i) + '.jpeg')
    i += 1

#with Image(filename='PB-raw0.jpeg') as top:
with Image(width=top.width,
           height=(int(0.75 * top.height) + (regular_page_height * (i-1))))  as stitch:
           stitch.composite(image=top, left=0, top=0)
           stitch.composite(image=Image(filename='PB-raw1.jpeg'), left=0, top=top.height)
           for j in range(2,i):
               stitch.composite(image=Image(filename='PB-raw' + str(j) + '.jpeg'), left=0, top=top.height + (regular_page_height * (j-1)))
           stitch.save(filename='PB-raw-stitch.jpeg')  #with Image(filename='PB-raw' + str(j) + '.jpeg') as bottom:


for k in range(0,i):
    os.remove('PB-raw' + str(k) + '.jpeg')

# Need to crop page one header information and then header entirely.

raw_stitch = Image(filename='PB-raw-stitch.jpeg').make_blob('JPEG')

txt = tool.image_to_string(
    PI.open(io.BytesIO(raw_stitch)),
    lang=lang,
    builder=pyocr.builders.TextBuilder()
)
txt = txt.replace('\n',' ')
txt = txt.replace('  ','\n')
for i in range(0,10):
    txt = txt.replace('\n\n','\n')
txt = txt.replace('\n','\n\n')
txt = txt.replace('&','a')
txt = txt.replace('1','I')

clean_up = 'bcdefghijklmnopqqrstuvwxyzABCDEFGHJKLMNOPQRSTUVWXYZ'
for i in range(0,len(clean_up)):
    txt = txt.replace(' ' + clean_up[i] + ' ', ' ' + clean_up[i])
    txt = txt.replace('.' + clean_up[i], '@' + clean_up[i])
    txt = txt.replace(clean_up[i] + ']', clean_up[i] + 'l')
    txt = txt.replace('a]','al')
    txt = txt.replace('a[','al')
    txt = txt.replace(clean_up[i] + '[', clean_up[i] + 'l')

txt = txt.replace(' ] ', ' I ')
txt = txt.replace(' [ ', ' I ')

numParagraphs = txt.count('\n\n') + 1

ocr_raw = Document()
ocr_raw.add_paragraph('')
ocr_raw.paragraphs[0].add_run(txt) # Each txt is generated per a page from the PDF
ocr_raw.save('OCR-raw.docx')

print(txt)
print('\n' + 'numParagraphs = ' + str(numParagraphs))
